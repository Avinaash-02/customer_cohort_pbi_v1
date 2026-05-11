import pandas as pd
import openpyxl
from docx import Document
import os

print("="*80)
print("WORKSPACE EXPLORATION - Customer Cohort & Retention Analysis")
print("="*80)

# 1. List all files in the directory
print("\n1. FILES IN THE DIRECTORY:")
print("-"*80)
for file in os.listdir("."):
    if os.path.isfile(file):
        print(f"  - {file}")

# 2. Analyze Excel file
print("\n2. EXCEL FILE STRUCTURE (NFX_Customer_Cohort_CaseStudy_Data.xlsx):")
print("-"*80)
excel_file = "NFX_Customer_Cohort_CaseStudy_Data.xlsx"
if os.path.exists(excel_file):
    wb = openpyxl.load_workbook(excel_file)
    sheets = wb.sheetnames
    print(f"Available Sheets: {sheets}")
    print("\n")
    
    # Read each sheet and display info
    for sheet in sheets:
        df = pd.read_excel(excel_file, sheet_name=sheet)
        print(f"Sheet: '{sheet}'")
        print(f"  Rows: {df.shape[0]}, Columns: {df.shape[1]}")
        print(f"  Columns: {list(df.columns)}")
        print(f"  Data Types:\n{df.dtypes}")
        print(f"  Preview (first 5 rows):")
        print(df.head())
        print(f"  Summary Statistics:")
        print(df.describe())
        print("\n")
else:
    print(f"Excel file not found: {excel_file}")

# 3. Analyze Word document
print("\n3. WORD DOCUMENT CONTENT (250411 Case Study Customer Cohort & Retention Analysis.docx):")
print("-"*80)
doc_file = "250411 Case Study Customer Cohort & Retention Analysis.docx"
if os.path.exists(doc_file):
    doc = Document(doc_file)
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    print("\nDocument Content (Key sections):")
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"  {para.text}")
else:
    print(f"Word document not found: {doc_file}")

print("\n" + "="*80)
